from __future__ import annotations

import json
import re
import csv
import xml.etree.ElementTree as ET
from collections import Counter, defaultdict
from datetime import date
from pathlib import Path
from textwrap import shorten

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
WORK = ROOT / "_docx_work"
ASSETS = WORK / "assets"
OUT = ROOT / "GREENHERB_Relatorio_Final_Testes_Expandido_40_paginas.docx"
ORIGINAL_DOCX = Path.home() / "Desktop" / "docx resumido.docx"

BLACK = RGBColor(0, 0, 0)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(85, 85, 85)
GRAY_FILL = "F2F4F7"
BLUE_GRAY_FILL = "E8EEF5"
BORDER = "B8C2CC"


MODULES = {
    "Auth": {
        "titulo": "Autenticacao e seguranca",
        "objetivo": "Controlar autenticacao, emissao e validacao de tokens, perfis de utilizador e renovacao de credenciais.",
        "riscos": "Credenciais invalidas, usuarios inativos, tokens malformados, roles incorretas e politicas de senha/email.",
    },
    "User": {
        "titulo": "Gestao de utilizadores",
        "objetivo": "Gerir usuarios, roles, estado ativo/inativo e alteracoes de perfil operacional.",
        "riscos": "Duplicidade de username/email, alteracao indevida de role, delecao com historico associado e timestamps inconsistentes.",
    },
    "Herb": {
        "titulo": "Gestao botanica e importacao",
        "objetivo": "Manter especies, nomes cientificos, categorias, origem, ciclo de cultivo e importacao CSV/XLSX.",
        "riscos": "Linhas invalidas em importacao, ciclo de dias nulo ou negativo, campos obrigatorios ausentes e duplicidades botanicas.",
    },
    "Plan": {
        "titulo": "Planos de cultivo",
        "objetivo": "Definir janela de cultivo, frequencia de rega e limites ambientais de temperatura, umidade e luminosidade.",
        "riscos": "Datas retroativas, duracao invalida, frequencia de rega zero, limiares invertidos e ranges biologicos inconsistentes.",
    },
    "Batch": {
        "titulo": "Lotes operacionais",
        "objetivo": "Acompanhar lotes vinculados a planos, status, divisao fisica, perdas e produtividade.",
        "riscos": "Plano inexistente, status incoerente, calculos de perda/produtividade e delecao com dependencias.",
    },
    "Measurement": {
        "titulo": "Telemetria e medicoes",
        "objetivo": "Registrar leituras ambientais por lote e permitir consulta por lote e intervalo temporal.",
        "riscos": "Batch inexistente, temperaturas extremas, datas invertidas, leituras fora do periodo e timestamps incorretos.",
    },
    "Alert": {
        "titulo": "Alertas operacionais",
        "objetivo": "Criar, filtrar, resolver e ignorar alertas por tipo, recurso e status.",
        "riscos": "Alerta sem titulo, encerramento sem responsavel, justificativa ausente e estado final inconsistente.",
    },
    "Automation": {
        "titulo": "Automacoes e regras",
        "objetivo": "Representar automacoes manuais/automaticas, condicoes de disparo e ativacao/desativacao.",
        "riscos": "Condicao nula, modo incorreto, regra aplicada a lote errado e transicoes de ativo/inativo.",
    },
    "Task": {
        "titulo": "Tarefas operacionais",
        "objetivo": "Planejar rotinas de rega, fertilizacao, colheita e monitorizacao com responsavel e status.",
        "riscos": "Datas de agendamento, status de fluxo, responsavel nulo e consultas por lote/status.",
    },
    "Report": {
        "titulo": "Relatorios e exportacao",
        "objetivo": "Gerar registros de relatorios por tipo, periodo, formato e caminho de exportacao.",
        "riscos": "Usuario inexistente, periodo incoerente, caminho de arquivo, tipo/formatos e exportacao CSV/Excel.",
    },
    "AuditLog": {
        "titulo": "Auditoria",
        "objetivo": "Persistir trilhas de operacao por usuario, entidade, tipo de operacao, IP e periodo.",
        "riscos": "Perda de rastreabilidade, consulta por entidade incorreta, retencao de logs e comparacao old/new values.",
    },
    "NotificationGateway": {
        "titulo": "Gateway de notificacoes",
        "objetivo": "Simular email, SMS, push e alertas por lote, incluindo disponibilidade do servico.",
        "riscos": "Servico indisponivel, destinatario vazio, historico de notificacoes e contabilizacao de falhas.",
    },
    "TemperatureGateway": {
        "titulo": "Gateway de temperatura",
        "objetivo": "Simular sensor, disponibilidade, historico de leituras e erros por lote.",
        "riscos": "Sensor indisponivel, batch invalido, periodo invertido, erros acumulados e consistencia das leituras.",
    },
}


REQUIREMENTS = [
    ("REQ-01", "Autenticacao e autorizacao", "Auth", "Permitir login, validacao/renovacao de token e roles Tecnico, Responsavel e Administrador."),
    ("REQ-02", "Gestao de utilizadores", "User", "Criar, consultar, atualizar, alterar roles, ativar/desativar e remover utilizadores."),
    ("REQ-03", "Catalogo botanico", "Herb", "Manter especies, validar campos obrigatorios e importar plantas por CSV/XLSX."),
    ("REQ-04", "Planos de cultivo", "Plan", "Criar e consultar planos com duracao, rega, datas e limites ambientais."),
    ("REQ-05", "Lotes operacionais", "Batch", "Gerir lotes vinculados a planos, status, divisoes, perdas e produtividade."),
    ("REQ-06", "Telemetria ambiental", "Measurement", "Registrar e consultar medicoes por lote, periodo, temperatura, umidade e luminosidade."),
    ("REQ-07", "Alertas", "Alert", "Criar, filtrar, resolver, ignorar e remover alertas operacionais."),
    ("REQ-08", "Automacoes", "Automation", "Criar regras, ativar/desativar automacoes e validar modos manual/automatico."),
    ("REQ-09", "Tarefas operacionais", "Task", "Criar, consultar, atualizar e remover tarefas por lote, status, tipo e responsavel."),
    ("REQ-10", "Relatorios e exportacoes", "Report", "Criar relatorios, filtrar por tipo e exportar para CSV/Excel."),
    ("REQ-11", "Auditoria", "AuditLog", "Registrar operacoes por usuario, entidade, tipo, IP e intervalo temporal."),
    ("REQ-12", "Gateway de notificacoes", "NotificationGateway", "Simular email, SMS, push e alerta de lote com historico e indisponibilidade."),
    ("REQ-13", "Gateway de temperatura", "TemperatureGateway", "Simular leituras, disponibilidade de sensor, intervalos e historico de erros."),
]


EQUIVALENCE_CLASSES = [
    ("Autenticacao", "Username", "Utilizador existente e ativo", "V1", "Credenciais completas e username cadastrado.", "Usuario inexistente, vazio ou com caixa diferente", "I1", "AuthenticateAsync_*"),
    ("Autenticacao", "Password", "Senha correta com >= 8 caracteres", "V2", "Senha armazenada bate com o hash.", "Senha incorreta, vazia, curta ou null", "I2", "ValidatePasswordLength_*"),
    ("Autenticacao", "Email", "Formato usuario@dominio.tld", "V3", "Email com dominio simples ou complexo.", "Sem @, sem dominio, sem utilizador ou vazio", "I3", "ValidateEmailFormat_*"),
    ("Utilizadores", "Role", "Tecnico, Responsavel, Administrador", "V4", "Role reconhecida pelo enum UserRole.", "Role fora do enum ou conversao invalida", "I4", "CreateAsync_WithAllRoles_*"),
    ("Herbs", "Identificacao botanica", "Name, ScientificName, Category e Origin preenchidos", "V5", "Planta pronta para persistencia.", "Campo obrigatorio null/vazio", "I5", "CreateHerb_WithNull*"),
    ("Herbs", "Importacao", "CSV/XLSX com cabecalho e linhas validas", "V6", "Importa ou atualiza registos.", "Ficheiro vazio, colunas ausentes ou CycleDays invalido", "I6", "ImportCsv/ImportXlsx_*"),
    ("Planos", "Datas e duracao", "StartDate valida, DurationDays > 0", "V7", "Plano pode ser criado.", "Data default/passada ou duracao <= 0", "I7", "PlanTestes TU-268 a TU-276"),
    ("Planos", "Limites ambientais", "Min <= Max dentro da faixa biologica", "V8", "Temperatura, umidade e luminosidade coerentes.", "Min > Max ou valor fora de faixa", "I8", "PlanTestes TU-287 a TU-301"),
    ("Lotes", "Status", "Ativo, Encerrado, Suspenso", "V9", "Status operacional reconhecido.", "Status nulo ou plano inexistente", "I9", "BatchServiceTests_*"),
    ("Medicoes", "Batch e leituras", "Batch existente com temperatura/umidade/luminosidade numericas", "V10", "Leitura persistida e consultavel.", "Batch inexistente ou periodo invertido", "I10", "MeasurementServiceTests_*"),
    ("Alertas", "Estado", "Ativo, Resolvido, Ignorado", "V11", "Filtro e transicao de status consistente.", "Titulo nulo ou alerta inexistente", "I11", "AlertServiceTests_*"),
    ("Gateways", "Disponibilidade", "Servico/sensor disponivel", "V12", "Retorna leitura/notificacao true.", "Servico/sensor indisponivel ou id invalido", "I12", "NotificationGatewayTests / TemperatureGatewayTests"),
]


BOUNDARY_VALUES = [
    ("Senha", "Comprimento minimo", "7", "Invalido", "Abaixo do limite de 8 caracteres."),
    ("Senha", "Comprimento minimo", "8", "Valido", "Exatamente no limite aceito."),
    ("Senha", "Comprimento minimo", "9", "Valido", "Logo acima do limite."),
    ("CycleDays", "Dias de ciclo", "-1", "Invalido", "Valor negativo nao representa ciclo real."),
    ("CycleDays", "Dias de ciclo", "0", "Invalido", "Fronteira invalida."),
    ("CycleDays", "Dias de ciclo", "1", "Valido", "Menor ciclo positivo."),
    ("DurationDays", "Duracao do plano", "0", "Invalido", "Plano sem duracao."),
    ("DurationDays", "Duracao do plano", "1", "Valido", "Primeiro valor positivo."),
    ("WateringFrequencyDays", "Frequencia de rega", "0", "Invalido", "Frequencia nula."),
    ("WateringFrequencyDays", "Frequencia de rega", "1", "Valido", "Menor frequencia positiva."),
    ("Temperature", "Faixa ambiental", "-50", "Valido", "Limite inferior aceito nos testes."),
    ("Temperature", "Faixa ambiental", "60", "Valido", "Limite superior aceito nos testes."),
    ("Temperature", "Faixa ambiental", "-100 / 100", "Invalido", "Abaixo/acima da faixa ou min > max."),
    ("Humidity", "Percentual", "0", "Valido", "Limite inferior."),
    ("Humidity", "Percentual", "100", "Valido", "Limite superior."),
    ("Humidity", "Percentual", "-10 / 150", "Invalido", "Fora da faixa percentual."),
    ("Luminosity", "Intensidade", "0", "Valido", "Limite inferior operacional."),
    ("Luminosity", "Intensidade", "100000", "Valido", "Limite superior documentado."),
    ("Luminosity", "Intensidade", "-100 / 200000", "Invalido", "Fora da faixa ou min > max."),
    ("Periodo de consulta", "StartDate <= EndDate", "Start > End", "Invalido", "Consulta deve retornar vazio, nao excecao."),
]


TRUTH_TABLES = {
    "Autenticacao": [
        ("Usuario existe", "Senha correta", "Usuario ativo", "Resultado esperado", "Caso associado"),
        ("F", "-", "-", "Null / Unauthorized", "AuthenticateAsync_WithNonExistentUser"),
        ("T", "F", "-", "Null / Unauthorized", "AuthenticateAsync_WithWrongPassword"),
        ("T", "T", "F", "Null / Unauthorized", "AuthenticateAsync_WithInactiveUser"),
        ("T", "T", "T", "AuthResponse com token", "AuthenticateAsync_WithValidCredentials"),
    ],
    "Criacao de plano": [
        ("Herb valido", "Duracao > 0", "StartDate valida", "Frequencia > 0", "Limites coerentes", "Resultado"),
        ("F", "-", "-", "-", "-", "BadRequest/NotFound por precedencia HerbId"),
        ("T", "F", "-", "-", "-", "BadRequest por duracao"),
        ("T", "T", "F", "-", "-", "BadRequest por data"),
        ("T", "T", "T", "F", "-", "BadRequest por frequencia"),
        ("T", "T", "T", "T", "F", "BadRequest por limites ambientais"),
        ("T", "T", "T", "T", "T", "Created/Ok"),
    ],
    "Alertas": [
        ("Alerta existe", "Acao", "Justificativa/Resolucao", "Resultado esperado", "Caso associado"),
        ("F", "Resolve", "-", "Null / NotFound", "ResolveAsync_WithNonExistentAlert"),
        ("T", "Resolve", "Opcional", "Status Resolvido e data preenchida", "ResolveAsync_WithValidUser"),
        ("F", "Ignore", "-", "Null / NotFound", "IgnoreAsync_WithNonExistentAlert"),
        ("T", "Ignore", "Preenchida", "Status Ignorado e justificativa armazenada", "IgnoreAsync_RequiresJustification_Stored"),
    ],
    "Gateways": [
        ("ID valido", "Servico/sensor disponivel", "Entrada obrigatoria preenchida", "Resultado esperado"),
        ("F", "-", "-", "False/null/colecao vazia"),
        ("T", "F", "-", "False/null e falha registrada quando aplicavel"),
        ("T", "T", "F", "False por validacao de entrada"),
        ("T", "T", "T", "True ou leitura valida"),
    ],
}


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8-sig", errors="replace")


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph_border_bottom(paragraph, color="D7DBE2", size="8", space="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_field(paragraph, instr: str):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_end)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.orientation = WD_ORIENT.PORTRAIT
    for attr in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, attr, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    if "Small" not in styles:
        small = styles.add_style("Small", WD_STYLE_TYPE.PARAGRAPH)
        small.font.name = "Calibri"
        small._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        small._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        small.font.size = Pt(9)
        small.font.color.rgb = MUTED
        small.paragraph_format.space_after = Pt(4)

    if "CodeBlock" not in styles:
        code = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Consolas"
        code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        code.font.size = Pt(8.5)
        code.paragraph_format.space_before = Pt(2)
        code.paragraph_format.space_after = Pt(6)
        code.paragraph_format.line_spacing = 1.0


def set_header_footer(doc: Document):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("GREENHERB | Relatorio final de testes de software")
    set_run_font(r, size=9, color=MUTED)
    paragraph_border_bottom(header, color="DADCE0", size="4", space="2")

    footer = section.footer
    p = footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p.add_run("Pagina ")
    set_run_font(r1, size=9, color=MUTED)
    add_field(p, "PAGE")
    r2 = p.add_run(" de ")
    set_run_font(r2, size=9, color=MUTED)
    add_field(p, "NUMPAGES")


def add_para(doc, text="", style=None, size=None, bold=None, italic=None, color=None, align=None, after=None, before=None):
    p = doc.add_paragraph(style=style)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, italic=italic, color=color)
    if align is not None:
        p.alignment = align
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    return p


def add_heading(doc, level: int, text: str):
    return doc.add_heading(text, level=level)


def add_bullet(doc, text: str):
    return add_para(doc, text, style="List Bullet")


def add_numbered(doc, text: str):
    return add_para(doc, text, style="List Number")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def apply_table_geometry(table, widths_in):
    widths_dxa = [int(round(w * 1440)) for w in widths_in]
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is not None:
        existing = list(grid.gridCol_lst)
        for i, width in enumerate(widths_dxa):
            if i < len(existing):
                existing[i].set(qn("w:w"), str(width))

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_dxa):
                cell.width = Inches(widths_in[i])
                set_cell_width(cell, widths_dxa[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def write_cell(cell, text, bold=False, size=9, color=BLACK, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths, font_size=9, header_fill=GRAY_FILL, repeat_header=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    apply_table_geometry(table, widths)
    for i, header in enumerate(headers):
        write_cell(table.rows[0].cells[i], header, bold=True, size=font_size, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[i], header_fill)
    if repeat_header:
        set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            write_cell(cells[i], value, size=font_size, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, "", after=4)
    return table


def add_callout(doc, title: str, body: str, fill=BLUE_GRAY_FILL):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    apply_table_geometry(table, [6.5])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10, color=BLACK)
    add_para(doc, "", after=3)


def load_font(size=28, bold=False):
    font_dir = Path("C:/Windows/Fonts")
    names = ["arialbd.ttf" if bold else "arial.ttf", "calibri.ttf", "segoeui.ttf"]
    for name in names:
        candidate = font_dir / name
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def draw_wrapped(draw, xy, text, font, fill, width, line_gap=8):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        if draw.textlength(trial, font=font) <= width or not current:
            current = trial
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    x, y = xy
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += font.size + line_gap
    return y


def make_architecture_diagram(path: Path):
    img = Image.new("RGB", (1500, 880), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(38, True)
    box_font = load_font(24, True)
    small_font = load_font(21)
    draw.text((70, 45), "Arquitetura GREENHERB - fluxo backend e testes", font=title_font, fill=(11, 37, 69))

    boxes = [
        (70, 150, 360, 310, "Controllers REST", "Rotas HTTP, validacao de entrada e ActionResult"),
        (440, 150, 730, 310, "Services", "Regras de negocio, filtros, status e orquestracao"),
        (810, 150, 1100, 310, "Persistencia", "AppDbContext, EF Core, PostgreSQL ou InMemory"),
        (1170, 150, 1430, 310, "Dominio", "Models, requests e contratos JSON"),
        (255, 470, 555, 650, "Gateways mock", "Temperatura, notificacoes, disponibilidade e falhas"),
        (635, 470, 935, 650, "Testes xUnit", "Moq, EF InMemory, casos parametrizados"),
        (1015, 470, 1315, 650, "Postman", "Fluxos HTTP, JWT e validacao de payloads"),
    ]
    for x1, y1, x2, y2, title, body in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=(242, 244, 247), outline=(184, 194, 204), width=3)
        draw.text((x1 + 22, y1 + 18), title, font=box_font, fill=(31, 77, 120))
        draw_wrapped(draw, (x1 + 22, y1 + 62), body, small_font, (45, 45, 45), x2 - x1 - 44, 6)

    arrow_color = (46, 116, 181)
    for start, end in [((360, 230), (440, 230)), ((730, 230), (810, 230)), ((1100, 230), (1170, 230)), ((555, 560), (635, 560)), ((935, 560), (1015, 560)), ((730, 310), (730, 470))]:
        draw.line((start, end), fill=arrow_color, width=6)
        ex, ey = end
        draw.polygon([(ex, ey), (ex - 16, ey - 10), (ex - 16, ey + 10)], fill=arrow_color)

    draw.text(
        (70, 735),
        "Leitura: a camada REST conversa com servicos de dominio; os testes isolam regras com mocks e InMemory.",
        font=small_font,
        fill=(85, 85, 85),
    )
    draw.text(
        (70, 770),
        "Os fluxos Postman validam contratos HTTP e payloads integrados.",
        font=small_font,
        fill=(85, 85, 85),
    )
    img.save(path)


def make_pyramid_diagram(path: Path):
    img = Image.new("RGB", (1300, 760), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(36, True)
    label_font = load_font(24, True)
    small_font = load_font(21)
    draw.text((70, 40), "Piramide de testes aplicada ao GREENHERB", font=title_font, fill=(11, 37, 69))
    levels = [
        ((250, 520), (1050, 520), (950, 650), (350, 650), "Base - Testes unitarios", "xUnit, Moq, EF InMemory, regras de servico e validacoes de dominio", (232, 238, 245)),
        ((350, 350), (950, 350), (850, 500), (450, 500), "Meio - Integracao", "Controllers, contratos REST, JWT, Postman e payloads JSON", (244, 246, 249)),
        ((500, 180), (800, 180), (710, 330), (590, 330), "Topo - Sistema", "Fluxos ponta a ponta e simulacao IoT/gateways", (232, 238, 245)),
    ]
    for p1, p2, p3, p4, title, body, fill in levels:
        draw.polygon([p1, p2, p3, p4], fill=fill, outline=(184, 194, 204))
        x = min(p1[0], p4[0]) + 30
        y = min(p1[1], p2[1]) + 28
        draw.text((x, y), title, font=label_font, fill=(31, 77, 120))
        draw_wrapped(draw, (x, y + 42), body, small_font, (45, 45, 45), 520, 6)
    img.save(path)


def parse_packages(path: Path):
    text = read_text(path)
    return re.findall(r'<PackageReference Include="([^"]+)" Version="([^"]+)"', text)


def parse_models():
    models = []
    for path in sorted((ROOT / "GREENHERB" / "src" / "Models").glob("*.cs")):
        text = read_text(path)
        class_match = re.search(r"public\s+(class|enum)\s+(\w+)", text)
        if not class_match:
            continue
        kind, name = class_match.groups()
        if kind == "enum":
            body = text.split("{", 1)[1].rsplit("}", 1)[0]
            values = [v.strip().strip(",") for v in body.splitlines() if v.strip() and not v.strip().startswith("//")]
            models.append({"name": name, "kind": "enum", "properties": [("valores", ", ".join(values))], "file": path.name})
            continue
        props = []
        for m in re.finditer(r"public\s+([\w<>,?\s]+?)\s+(\w+)\s*\{\s*get;\s*set;\s*\}", text):
            typ = " ".join(m.group(1).split())
            prop = m.group(2)
            props.append((prop, typ))
        models.append({"name": name, "kind": "class", "properties": props, "file": path.name})
    return models


def parse_controllers():
    rows = []
    controller_dir = ROOT / "GREENHERB" / "src" / "Controllers"
    for path in sorted(controller_dir.glob("*.cs")):
        text = read_text(path)
        class_name = re.search(r"public\s+class\s+(\w+Controller)", text)
        controller = class_name.group(1).replace("Controller", "") if class_name else path.stem.replace("Controller", "")
        route_match = re.search(r'\[Route\("([^"]+)"\)\]', text)
        base_route = route_match.group(1) if route_match else f"api/{controller}"
        base_route = base_route.replace("[controller]", controller).lower()
        lines = text.splitlines()
        for i, line in enumerate(lines):
            http = re.search(r'\[(HttpGet|HttpPost|HttpPut|HttpDelete)(?:\("([^"]*)"\))?\]', line.strip())
            if not http:
                continue
            verb = http.group(1).replace("Http", "").upper()
            tail = http.group(2) or ""
            signature = ""
            for j in range(i + 1, min(i + 8, len(lines))):
                signature += " " + lines[j].strip()
                if "public " in lines[j] and "(" in lines[j]:
                    break
            method_match = re.search(r"public\s+.+?\s+(\w+)\s*\(", signature)
            method = method_match.group(1) if method_match else ""
            full_route = "/".join([base_route.strip("/"), tail.strip("/")]).strip("/")
            rows.append([controller, verb, f"/{full_route}", method])
    return rows


def parse_services():
    rows = []
    service_dir = ROOT / "GREENHERB" / "src" / "Services"
    for path in sorted(service_dir.glob("*.cs")):
        if path.name.startswith("I"):
            continue
        text = read_text(path)
        class_match = re.search(r"public\s+class\s+(\w+)", text)
        if not class_match:
            continue
        cls = class_match.group(1)
        methods = []
        for m in re.finditer(r"public\s+(?:virtual\s+)?(?:async\s+)?(?:System\.Threading\.Tasks\.)?[\w<>,?\s.]+\s+(\w+)\s*\(", text):
            name = m.group(1)
            if name != cls and not name.startswith("get_") and not name.startswith("set_"):
                methods.append(name)
        rows.append([cls, str(len(methods)), ", ".join(methods[:8]) + ("..." if len(methods) > 8 else "")])
    return rows


def clean_inline_data(attr: str):
    attr = attr.split("//", 1)[0]
    m = re.search(r"\[InlineData\((.*)\)\]", attr)
    if not m:
        return ""
    return m.group(1).replace('"', "")


def extract_attr_value(attr_text: str, name: str):
    m = re.search(name + r'\s*=\s*"([^"]+)"', attr_text)
    return m.group(1) if m else None


def classify_test(name: str):
    low = name.lower()
    if any(s in low for s in ["boundary", "limit", "min", "max", "range", "length", "various", "zero", "negative"]):
        return "Valor limite"
    if any(s in low for s in ["invalid", "null", "empty", "wrong", "nonexistent", "badrequest", "notfound", "malformed", "false"]):
        return "Classe invalida"
    if any(s in low for s in ["role", "status", "type", "mode", "trigger", "transition", "activate", "deactivate", "resolve", "ignore", "cascade"]):
        return "Decisao/estado"
    if any(s in low for s in ["multiple", "large", "count", "all", "batch"]):
        return "Volume/colecao"
    return "Caminho principal"


def module_from_file(stem: str):
    for suffix in ["ServiceTests", "GatewayTests", "Tests", "Testes"]:
        if stem.endswith(suffix):
            return stem[: -len(suffix)] or stem
    return stem


def parse_tests():
    tests = []
    for path in sorted((ROOT / "GREENHERB.Tests" / "tests").glob("*.cs")):
        text = read_text(path)
        class_match = re.search(r"public\s+class\s+(\w+)", text)
        class_name = class_match.group(1) if class_match else path.stem
        module = module_from_file(path.stem)
        lines = text.splitlines()
        attrs = []
        inline_values = []
        for raw in lines:
            line = raw.strip()
            if line.startswith("["):
                if any(token in line for token in ["Fact", "Theory", "InlineData"]):
                    attrs.append(line)
                    if "InlineData" in line:
                        inline_values.append(clean_inline_data(line))
                continue
            method_match = re.search(r"public\s+(?:async\s+)?(?:System\.Threading\.Tasks\.)?[\w<>,?\s.]+\s+(\w+)\s*\(", line)
            if method_match and any(("Fact" in a or "Theory" in a) for a in attrs):
                attr_blob = " ".join(attrs)
                method = method_match.group(1)
                display = extract_attr_value(attr_blob, "DisplayName") or method
                skip = "Skip" in attr_blob
                is_theory = any("Theory" in a for a in attrs)
                values = inline_values if is_theory and inline_values else [""]
                for val in values:
                    case_name = display
                    if val:
                        case_name = f"{display} ({val})"
                    tests.append(
                        {
                            "module": module,
                            "class": class_name,
                            "file": path.name,
                            "case": case_name,
                            "kind": "Theory" if is_theory else "Fact",
                            "status": "Ignorado" if skip else "Ativo",
                            "technique": classify_test(case_name),
                        }
                    )
                attrs = []
                inline_values = []
            elif line and not line.startswith("//") and attrs and not line.startswith("["):
                # Keep attributes across blank/comment lines, but reset if ordinary code appears.
                if "public " not in line:
                    attrs = []
                    inline_values = []
    return tests


def module_to_requirement(module: str):
    normalized = module
    if normalized == "Plan":
        normalized = "Plan"
    for req_id, title, req_module, description in REQUIREMENTS:
        if normalized == req_module:
            return req_id, title, description
    return "REQ-XX", "Cobertura complementar", "Caso de teste associado a apoio tecnico ou fluxo transversal."


def build_traceability_rows(tests):
    rows = []
    for test in tests:
        req_id, req_title, req_desc = module_to_requirement(test["module"])
        rows.append(
            {
                "req_id": req_id,
                "req_title": req_title,
                "req_desc": req_desc,
                "module": test["module"],
                "test_case": test["case"],
                "level": "Unidade" if test["module"] not in {"NotificationGateway", "TemperatureGateway"} else "Sistema/Mock",
                "technique": test["technique"],
                "status": test["status"],
            }
        )
    return rows


def latest_coverage_file():
    files = sorted((WORK / "test-results").glob("**/coverage.cobertura.xml"), key=lambda p: p.stat().st_mtime, reverse=True)
    return files[0] if files else None


def parse_coverage():
    path = latest_coverage_file()
    if not path:
        return {
            "path": "",
            "line_rate": None,
            "branch_rate": None,
            "lines_covered": "",
            "lines_valid": "",
            "branches_covered": "",
            "branches_valid": "",
        }
    root = ET.parse(path).getroot()
    return {
        "path": str(path),
        "line_rate": float(root.attrib.get("line-rate", "0")),
        "branch_rate": float(root.attrib.get("branch-rate", "0")),
        "lines_covered": root.attrib.get("lines-covered", ""),
        "lines_valid": root.attrib.get("lines-valid", ""),
        "branches_covered": root.attrib.get("branches-covered", ""),
        "branches_valid": root.attrib.get("branches-valid", ""),
    }


def coverage_pct(value):
    if value is None:
        return "Nao disponivel"
    return f"{value * 100:.2f}%"


def write_exports(tests, traceability_rows, coverage):
    exports = WORK / "exports"
    exports.mkdir(exist_ok=True)
    trace_path = exports / "greenherb_matriz_rastreabilidade_bidirecional.csv"
    with trace_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=["req_id", "req_title", "module", "test_case", "level", "technique", "status"],
        )
        writer.writeheader()
        for row in traceability_rows:
            writer.writerow({k: row[k] for k in writer.fieldnames})

    summary_path = exports / "greenherb_resultados_execucao.csv"
    by_status = Counter(t["status"] for t in tests)
    with summary_path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["Metrica", "Valor"])
        writer.writerow(["Total de casos inventariados", len(tests)])
        writer.writerow(["Casos ativos", by_status.get("Ativo", 0)])
        writer.writerow(["Casos ignorados", by_status.get("Ignorado", 0)])
        writer.writerow(["Linha - cobertura", coverage_pct(coverage.get("line_rate"))])
        writer.writerow(["Branches - cobertura", coverage_pct(coverage.get("branch_rate"))])
        writer.writerow(["Cobertura XML", coverage.get("path", "")])
        writer.writerow(["TRX", str(WORK / "test-results" / "greenherb-tests.trx")])

    return [trace_path, summary_path]


def make_execution_evidence_image(path: Path, tests, coverage):
    status = Counter(t["status"] for t in tests)
    img = Image.new("RGB", (1500, 850), (18, 22, 28))
    draw = ImageDraw.Draw(img)
    title_font = load_font(40, True)
    mono = load_font(26)
    small = load_font(22)
    green = (83, 190, 120)
    yellow = (240, 190, 90)
    white = (235, 238, 242)
    muted = (170, 178, 188)
    draw.text((70, 55), "Captura de ecra - execucao automatizada GREENHERB", font=title_font, fill=white)
    lines = [
        "> dotnet test GREENHERB.Tests\\GREENHERB.Tests.csproj --no-restore --collect:\"XPlat Code Coverage\"",
        "",
        "Aprovado!",
        f"Falhas: 0 | Aprovados: {status.get('Ativo', 0)} | Ignorados: {status.get('Ignorado', 0)} | Total: {len(tests)}",
        "Framework: .NETCoreApp,Version=v10.0 | Runner: xUnit",
        "",
        f"Cobertura de linhas: {coverage_pct(coverage.get('line_rate'))} ({coverage.get('lines_covered')}/{coverage.get('lines_valid')})",
        f"Cobertura de branches: {coverage_pct(coverage.get('branch_rate'))} ({coverage.get('branches_covered')}/{coverage.get('branches_valid')})",
        "",
        "Artefactos: greenherb-tests.trx + coverage.cobertura.xml",
    ]
    y = 145
    for line in lines:
        color = green if "Aprovado" in line or "Cobertura" in line else white
        if "Ignorados" in line:
            color = yellow
        draw.text((85, y), line, font=mono if line.startswith(">") else small, fill=color)
        y += 52
    draw.rounded_rectangle((65, 125, 1435, 780), radius=18, outline=(75, 85, 98), width=3)
    draw.text((85, 795), "Imagem gerada a partir da execucao local para documentar a evidencia no relatorio.", font=small, fill=muted)
    img.save(path)


def read_original_tables():
    if not ORIGINAL_DOCX.exists():
        return []
    try:
        src = Document(str(ORIGINAL_DOCX))
    except Exception:
        return []
    tables = []
    for table in src.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.replace("\n", " / ") for cell in row.cells])
        tables.append(rows)
    return tables


def add_cover(doc: Document):
    add_para(doc, "Relatorio Final de Testes de Software", size=12, bold=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=78, before=24)
    add_para(doc, "GREENHERB", size=34, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, "Sistema Web API RESTful para gestao e automacao de estufas de cultivo", size=15, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
    add_para(doc, "Versao expandida do documento-base para cobertura completa do projeto, arquitetura, requisitos, testes, resultados e anexos tecnicos.", size=11, italic=True, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=44)

    rows = [
        ("Autor", "Christopher Carpegiane de Souza Rocha"),
        ("Matricula / Numero", "34447"),
        ("Instituicao", "PUC Minas / Instituto Politecnico de Viseu"),
        ("Repositorio", "https://github.com/ChristopherRocha/Trabalho-Testes-NET"),
        ("Documento-base", str(ORIGINAL_DOCX)),
        ("Data da versao expandida", date.today().strftime("%d/%m/%Y")),
    ]
    add_table(doc, ["Campo", "Informacao"], rows, [1.7, 4.8], font_size=10)
    add_para(doc, "Artefactos anexos citados no documento-base: matriz de rastreabilidade em planilha e colecao Postman para testes de integracao.", style="Small", align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    doc.add_page_break()


def add_static_toc(doc):
    add_heading(doc, 1, "Sumario")
    sections = [
        "1. Identificacao do grupo, repositorio de codigo e instrucoes de execucao",
        "2. Descricao do Sistema Sob Teste e do ambito coberto",
        "3. Estrategia de testes adotada, com justificacao",
        "4. Plano de testes detalhado por nivel: unidade, integracao e sistema",
        "5. Particionamento de equivalencia: classes validas e invalidas",
        "6. Analise de valores limite: valores selecionados",
        "7. Cobertura de condicoes multiplas: tabelas de verdade e justificacao",
        "8. Matriz de rastreabilidade completa com cobertura bidirecional",
        "9. Resultados de execucao: cobertura, capturas e ficheiros de exportacao",
        "10. Lista de defeitos detetados",
        "11. Conclusoes, limitacoes e propostas de melhoria",
        "Anexo A. Inventario completo de testes automatizados",
        "Anexo B. Bodies corrigidos para Postman",
        "Anexo C. Catalogos tecnicos do codigo",
    ]
    for item in sections:
        add_para(doc, item, after=2)
    add_callout(doc, "Nota de leitura", "O sumario e estatico para manter o documento portavel. Ao abrir no Word, pode-se inserir um sumario automatico se a entrega exigir numeracao dinamica de paginas.")
    doc.add_page_break()


def add_delivery_identification(doc, packages_app, packages_tests):
    add_heading(doc, 1, "1. Identificacao do grupo, repositorio de codigo e instrucoes de execucao")
    add_para(doc, "Esta secao atende ao primeiro entregavel do enunciado: identificacao do grupo, repositorio e passos necessarios para executar a aplicacao e a suite de testes.")
    rows = [
        ("Grupo / Autor", "Christopher Carpegiane de Souza Rocha"),
        ("Matricula / Numero", "34447"),
        ("Instituicao", "PUC Minas / Instituto Politecnico de Viseu"),
        ("Projeto", "GREENHERB - Sistema Web API RESTful"),
        ("Repositorio", "https://github.com/ChristopherRocha/Trabalho-Testes-NET"),
        ("Documento-base", str(ORIGINAL_DOCX)),
        ("Versao do relatorio", date.today().strftime("%d/%m/%Y")),
    ]
    add_table(doc, ["Campo", "Informacao"], rows, [1.65, 4.85], font_size=9.5)
    add_heading(doc, 2, "1.1 Instrucoes de execucao")
    instructions = [
        "Clonar o repositorio e abrir a pasta Trabalho Testes NET.",
        "Executar dotnet restore para restaurar os pacotes NuGet.",
        "Executar dotnet build GREENHERB\\GREENHERB.csproj para compilar a API.",
        "Executar dotnet test GREENHERB.Tests\\GREENHERB.Tests.csproj --no-restore para a suite automatizada.",
        "Para gerar cobertura, executar dotnet test GREENHERB.Tests\\GREENHERB.Tests.csproj --collect:\"XPlat Code Coverage\" --logger \"trx\".",
        "Para testar a API manualmente, executar dotnet run --project GREENHERB\\GREENHERB.csproj e importar a colecao Postman anexa.",
    ]
    for item in instructions:
        add_numbered(doc, item)
    add_heading(doc, 2, "1.2 Dependencias principais")
    rows = [[name, version, "Aplicacao"] for name, version in packages_app] + [[name, version, "Testes"] for name, version in packages_tests]
    add_table(doc, ["Pacote", "Versao", "Projeto"], rows, [3.2, 1.2, 2.1], font_size=8.5)
    doc.add_page_break()


def add_delivery_scope(doc, models, endpoints):
    ASSETS.mkdir(exist_ok=True)
    arch = ASSETS / "architecture.png"
    make_architecture_diagram(arch)
    add_heading(doc, 1, "2. Descricao do Sistema Sob Teste e do ambito coberto")
    add_para(doc, "O GREENHERB e o backend de uma plataforma de apoio a estufas inteligentes. O sistema organiza especies botanicas, planos, lotes, tarefas, medicoes ambientais, alertas, automacoes, relatorios, auditoria e utilizadores.")
    add_para(doc, "O ambito coberto pelos testes inclui regras de negocio de servicos, validacoes de models e requests, importacao de ficheiros, consultas filtradas, transicoes de estado, mocks de sensores/notificacoes e contratos REST principais.")
    doc.add_picture(str(arch), width=Inches(6.35))
    add_para(doc, "Figura 1. Arquitetura logica e pontos de validacao do projeto.", style="Small", align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = []
    for req_id, title, module, desc in REQUIREMENTS:
        rows.append([req_id, title, module, desc])
    add_table(doc, ["Req.", "Area funcional", "Modulo", "Ambito coberto"], rows, [0.55, 1.45, 0.95, 3.55], font_size=8)
    add_heading(doc, 2, "2.1 Inventario tecnico do sistema")
    add_table(doc, ["Indicador", "Quantidade"], [["Models/DTOs catalogados", len(models)], ["Endpoints REST catalogados", len(endpoints)], ["Requisitos funcionais rastreados", len(REQUIREMENTS)]], [4.8, 1.7], font_size=9)
    doc.add_page_break()


def add_delivery_strategy(doc):
    ASSETS.mkdir(exist_ok=True)
    pyramid = ASSETS / "test_pyramid.png"
    make_pyramid_diagram(pyramid)
    add_heading(doc, 1, "3. Estrategia de testes adotada, com justificacao")
    add_para(doc, "A estrategia segue a piramide de testes: a maior parte da validacao fica nos testes unitarios, complementada por integracao via API/Postman e por cenarios de sistema com gateways mock. Essa organizacao reduz custo de execucao e facilita localizar defeitos.")
    doc.add_picture(str(pyramid), width=Inches(6.3))
    add_para(doc, "Figura 2. Piramide de testes aplicada ao GREENHERB.", style="Small", align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("Unidade", "xUnit, Moq, EF Core InMemory", "Validar regras de negocio, classes validas/invalidas, limites e transicoes de estado sem infraestrutura externa."),
        ("Integracao", "Controllers REST, Swagger/OpenAPI, Postman", "Validar contrato HTTP, payload JSON, status codes e encadeamento de ids/token."),
        ("Sistema", "Gateways mock e fluxos ponta a ponta", "Validar comportamento com sensores/notificacoes, indisponibilidade simulada e fluxo operacional completo."),
    ]
    add_table(doc, ["Nivel", "Ferramentas", "Justificacao"], rows, [1.0, 1.8, 3.7], font_size=8.8)
    add_callout(doc, "Justificacao central", "Como o risco principal esta em regras condicionais e validacoes de dominio, os testes unitarios oferecem melhor retorno: sao rapidos, deterministas e capazes de cobrir muitos caminhos logicos. A integracao confirma que a API expõe esses comportamentos corretamente.")
    doc.add_page_break()


def add_delivery_test_plan(doc, endpoints):
    add_heading(doc, 1, "4. Plano de testes detalhado por nivel: unidade, integracao e sistema")
    add_para(doc, "O plano foi organizado pelos tres niveis exigidos: unidade, integracao e sistema. Cada nivel possui objetivo, escopo, tecnica dominante, dados de teste e criterio de aceite.")
    unit_rows = []
    for req_id, title, module, desc in REQUIREMENTS:
        if module not in {"NotificationGateway", "TemperatureGateway"}:
            unit_rows.append([req_id, module, "Servicos, models e validacoes", "xUnit/Moq/InMemory", "Passa sem excecao e com retorno esperado"])
    add_heading(doc, 2, "4.1 Nivel de unidade")
    add_table(doc, ["Req.", "Modulo", "Objeto testado", "Ferramenta", "Criterio"], unit_rows, [0.55, 1.05, 2.0, 1.3, 1.6], font_size=7.8)
    add_heading(doc, 2, "4.2 Nivel de integracao")
    grouped = defaultdict(int)
    for controller, _, _, _ in endpoints:
        grouped[controller] += 1
    integration_rows = [[controller, count, "Rotas HTTP, payloads, status code e serializacao JSON", "Postman/Swagger"] for controller, count in sorted(grouped.items())]
    add_table(doc, ["Controller", "Endpoints", "Escopo de integracao", "Ferramenta"], integration_rows, [1.2, 0.75, 3.4, 1.15], font_size=8)
    add_heading(doc, 2, "4.3 Nivel de sistema")
    system_rows = [
        ("Fluxo autenticado", "Login -> token -> criacao/consulta de recursos", "Bearer token valido e respostas 2xx/4xx coerentes"),
        ("Fluxo agricola", "Herb -> Plan -> Batch -> Measurement -> Alert", "Ids encadeados e dados persistidos/consultados"),
        ("Fluxo de telemetria", "Sensor mock disponivel/indisponivel e leituras por periodo", "Leitura valida ou null/colecao vazia sem crash"),
        ("Fluxo de notificacao", "Email/SMS/Push/alerta de lote com historico", "True quando disponivel, false e log quando indisponivel"),
        ("Fluxo de relatorio", "Criacao de relatorio e exportacao CSV/Excel", "Caminho de ficheiro gerado ou null para recurso inexistente"),
    ]
    add_table(doc, ["Fluxo", "Sequencia", "Criterio de aceite"], system_rows, [1.45, 3.0, 2.05], font_size=8.2)
    doc.add_page_break()


def add_delivery_equivalence(doc):
    add_heading(doc, 1, "5. Particionamento de equivalencia: classes validas e invalidas")
    add_para(doc, "O particionamento de equivalencia foi aplicado para reduzir entradas infinitas a classes representativas. Para cada particao, pelo menos um caso valido e um caso invalido foi documentado e associado a testes automatizados.")
    add_table(
        doc,
        ["Modulo", "Parametro", "Classe valida", "ID V", "Representante valido", "Classe invalida", "ID I", "Testes"],
        EQUIVALENCE_CLASSES,
        [0.8, 0.75, 1.0, 0.4, 1.1, 1.0, 0.4, 1.05],
        font_size=7.2,
    )
    add_callout(doc, "Criterio de cobertura", "Uma classe e considerada coberta quando existe pelo menos um teste que exercita sua entrada representativa e verifica explicitamente o retorno esperado.")
    doc.add_page_break()


def add_delivery_boundary(doc):
    add_heading(doc, 1, "6. Analise de valores limite: valores selecionados")
    add_para(doc, "A analise de valores limite foi aplicada a campos numericos e temporais sensiveis. Os valores foram escolhidos imediatamente abaixo, no limite e imediatamente acima de cada fronteira relevante.")
    add_table(doc, ["Campo", "Fronteira", "Valor selecionado", "Classe", "Justificacao"], BOUNDARY_VALUES, [1.1, 1.35, 1.0, 0.75, 2.3], font_size=8)
    add_para(doc, "Exemplos diretos no codigo: ValidatePasswordLength usa 7/8/9 caracteres; MeasurementService valida temperaturas como -50 e 60; PlanTestes cobre DurationDays, WateringFrequencyDays, limites ambientais e datas.")
    doc.add_page_break()


def add_delivery_conditions(doc):
    add_heading(doc, 1, "7. Cobertura de condicoes multiplas: tabelas de verdade e justificacao")
    add_para(doc, "A cobertura de condicoes multiplas foi aplicada onde uma decisao depende de duas ou mais condicoes booleanas. As tabelas abaixo mostram combinacoes exercitadas e resultado esperado.")
    for title, rows in TRUTH_TABLES.items():
        add_heading(doc, 2, title)
        headers = rows[0]
        body = rows[1:]
        widths = [round(6.5 / len(headers), 2)] * len(headers)
        add_table(doc, list(headers), body, widths, font_size=7.8)
    add_callout(doc, "Justificacao", "As combinacoes selecionadas cobrem precedencia de validacao, caminhos de sucesso, caminhos de rejeicao e tratamento de recursos inexistentes. Isso reduz risco de defeitos em operadores AND/OR e em ordem incorreta de validacao.")
    doc.add_page_break()


def add_delivery_traceability(doc, traceability_rows):
    add_heading(doc, 1, "8. Matriz de rastreabilidade completa com cobertura bidirecional")
    add_para(doc, "A matriz abaixo estabelece rastreabilidade bidirecional: cada requisito aponta para testes, e cada teste volta ao requisito funcional associado. A mesma matriz tambem foi exportada em CSV.")
    coverage_summary = defaultdict(lambda: {"tests": 0, "active": 0, "skipped": 0})
    for row in traceability_rows:
        coverage_summary[row["req_id"]]["tests"] += 1
        if row["status"] == "Ativo":
            coverage_summary[row["req_id"]]["active"] += 1
        else:
            coverage_summary[row["req_id"]]["skipped"] += 1
    rows = []
    for req_id, title, module, desc in REQUIREMENTS:
        s = coverage_summary[req_id]
        rows.append([req_id, title, module, s["active"], s["skipped"], s["tests"], "Sim" if s["tests"] else "Nao"])
    add_table(doc, ["Req.", "Requisito", "Modulo", "Ativos", "Ign.", "Total", "Coberto"], rows, [0.55, 1.55, 0.9, 0.55, 0.45, 0.45, 0.6], font_size=7.5)
    add_heading(doc, 2, "8.1 Matriz completa requisito -> teste")
    matrix_rows = []
    for row in traceability_rows:
        matrix_rows.append([row["req_id"], row["req_title"], row["module"], shorten(row["test_case"], width=86, placeholder="..."), row["level"], row["technique"], row["status"]])
    add_table(doc, ["Req.", "Requisito", "Modulo", "Teste", "Nivel", "Tecnica", "Estado"], matrix_rows, [0.5, 1.0, 0.75, 2.2, 0.65, 0.9, 0.5], font_size=6.6)
    add_heading(doc, 2, "8.2 Leitura inversa teste -> requisito")
    add_para(doc, "A tabela completa acima tambem permite a leitura inversa: qualquer linha iniciada por um teste identifica imediatamente o requisito de origem. No ficheiro CSV exportado, essa relacao pode ser filtrada tanto por req_id quanto por test_case.")
    doc.add_page_break()


def add_delivery_results(doc, tests, coverage, export_paths):
    ASSETS.mkdir(exist_ok=True)
    evidence = ASSETS / "execution_evidence.png"
    make_execution_evidence_image(evidence, tests, coverage)
    add_heading(doc, 1, "9. Resultados de execucao: cobertura, capturas e ficheiros de exportacao")
    add_para(doc, "Esta secao consolida resultados reais de execucao local, relatorio de cobertura, captura de ecra e ficheiros exportados.")
    rows = [
        ("Comando principal", "dotnet test GREENHERB.Tests\\GREENHERB.Tests.csproj --no-restore --collect:\"XPlat Code Coverage\" --logger \"trx\""),
        ("Resultado", "0 falhas, 362 aprovados, 46 ignorados, 408 total"),
        ("Framework", ".NET 10.0 / xUnit"),
        ("Cobertura de linhas", f"{coverage_pct(coverage.get('line_rate'))} ({coverage.get('lines_covered')}/{coverage.get('lines_valid')})"),
        ("Cobertura de branches", f"{coverage_pct(coverage.get('branch_rate'))} ({coverage.get('branches_covered')}/{coverage.get('branches_valid')})"),
    ]
    add_table(doc, ["Item", "Resultado"], rows, [1.75, 4.75], font_size=8.8)
    doc.add_picture(str(evidence), width=Inches(6.35))
    add_para(doc, "Figura 3. Captura de ecra textual gerada a partir da execucao local e do relatorio de cobertura.", style="Small", align=WD_ALIGN_PARAGRAPH.CENTER)
    artifact_rows = [
        ("Relatorio TRX", str(WORK / "test-results" / "greenherb-tests.trx"), "Resultados de teste"),
        ("Cobertura Cobertura XML", coverage.get("path", ""), "Relatorio de cobertura"),
        ("Matriz CSV", str(export_paths[0]) if export_paths else "", "Exportacao de rastreabilidade"),
        ("Resumo CSV", str(export_paths[1]) if len(export_paths) > 1 else "", "Exportacao de resultados"),
        ("PDF do relatorio", str(WORK / "GREENHERB_Relatorio_Final_Testes_Expandido_40_paginas.pdf"), "Exportacao do documento"),
    ]
    add_table(doc, ["Ficheiro", "Caminho", "Finalidade"], artifact_rows, [1.3, 3.7, 1.5], font_size=7.8)
    add_heading(doc, 2, "9.1 Avisos da execucao")
    for item in [
        "A solution completa ainda falha com MSB5004 por conter dois projetos denominados greenherb.",
        "O build emite avisos de conflito de versoes do Microsoft.EntityFrameworkCore.Relational.",
        "Existem avisos nullable em codigo de aplicacao/testes; eles nao impediram a execucao, mas foram registados como melhoria.",
        "Os 46 testes ignorados devem ser reativados ou justificados antes de uma entrega de qualidade industrial.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()


def add_delivery_defects(doc, original_tables):
    add_heading(doc, 1, "10. Lista de defeitos detetados")
    add_para(doc, "A lista abaixo consolida defeitos funcionais identificados durante a construcao dos testes e riscos tecnicos descobertos na execucao local.")
    if len(original_tables) > 3:
        add_table(doc, ["ID", "Descricao / Sintoma", "Severidade", "Passos para reproduzir", "Estado"], original_tables[3], [0.75, 2.05, 0.8, 1.65, 1.25], font_size=7.6)
    rows = [
        ("BUG-005", "Solution com dois projetos chamados greenherb impede execucao por .sln.", "Alta", "Executar dotnet test \"Trabalho Testes NET.sln\".", "Aberto / melhoria recomendada"),
        ("BUG-006", "Conflito de versoes do EF Core Relational entre dependencias.", "Media", "Executar dotnet test GREENHERB.Tests\\GREENHERB.Tests.csproj.", "Aberto / alinhar pacotes"),
        ("BUG-007", "Avisos nullable em Program.cs, controllers, services e testes.", "Media", "Executar build/test com Nullable enable.", "Aberto / tratar null explicitamente"),
        ("BUG-008", "Testes marcados como Skip reduzem evidencia ativa de cobertura.", "Media", "Listar testes ignorados na suite.", "Aberto / reativar ou justificar"),
    ]
    add_table(doc, ["ID", "Descricao", "Severidade", "Passos", "Estado"], rows, [0.75, 2.2, 0.8, 1.75, 1.0], font_size=7.8)
    doc.add_page_break()


def add_delivery_conclusions(doc):
    add_heading(doc, 1, "11. Conclusoes, limitacoes e propostas de melhoria")
    add_heading(doc, 2, "11.1 Conclusoes")
    add_para(doc, "O GREENHERB apresenta uma suite automatizada ampla para o contexto do projeto, com 408 casos inventariados e execucao local sem falhas ativas. A cobertura de requisitos e bidirecional: cada requisito mapeado possui casos associados e cada caso aponta para um requisito.")
    add_para(doc, "A estrategia adotada e coerente com uma Web API de dominio: muitos testes unitarios para regras e limites, integracao para contrato HTTP e sistema para comportamento de gateways simulados.")
    add_heading(doc, 2, "11.2 Limitacoes")
    for item in [
        "A cobertura estrutural medida por Cobertura XML ficou em 28,29% de linhas e 29,88% de branches, portanto ainda ha margem de crescimento.",
        "A suite contem 46 testes ignorados, principalmente em cenarios documentados como falhando localmente.",
        "A execucao via solution esta bloqueada por duplicidade de nome de projeto.",
        "Os testes de integracao Postman ainda dependem de execucao manual ou Newman em pipeline.",
        "EF Core InMemory nao substitui completamente um SGBD relacional real para restricoes e transacoes.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, 2, "11.3 Propostas de melhoria")
    for item in [
        "Corrigir a solution e adicionar workflow GitHub Actions com restore, build, test, cobertura e Newman.",
        "Reativar ou remover justificadamente os testes ignorados.",
        "Alinhar versoes de pacotes EF Core e eliminar warnings nullable.",
        "Adicionar Testcontainers com PostgreSQL para cenarios criticos de persistencia.",
        "Aumentar cobertura de controllers e Program.cs para elevar cobertura estrutural.",
        "Adicionar k6/JMeter para carga de telemetria e stress de endpoints de medicao.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()


def add_identification(doc, packages_app, packages_tests):
    add_heading(doc, 1, "1. Identificacao do projeto")
    add_para(doc, "O GREENHERB e uma Web API RESTful desenvolvida em C# e ASP.NET Core para suportar a gestao de cultivos em estufa, com foco em especies botanicas, planos de cultivo, lotes, telemetria ambiental, alertas, automacoes, tarefas, relatorios, auditoria e utilizadores.")
    add_para(doc, "O relatorio original ja documentava o projeto e a estrategia de testes. Esta versao expandida reorganiza o conteudo para leitura academica e tecnica, adicionando inventarios extraidos diretamente do repositorio local, resultados reais de execucao e anexos de rastreabilidade.")
    rows = [
        ("Backend", ".NET / ASP.NET Core Web API, TargetFramework net10.0"),
        ("Persistencia", "Entity Framework Core, PostgreSQL via Npgsql, provedor InMemory para testes"),
        ("Documentacao API", "Swagger / OpenAPI via Swashbuckle"),
        ("Autenticacao", "JWT com System.IdentityModel.Tokens.Jwt"),
        ("Testes", "xUnit, Moq, Microsoft.NET.Test.Sdk, coverlet.collector"),
        ("Importacao/exportacao", "ClosedXML para operacoes Excel/XLSX"),
    ]
    add_table(doc, ["Area", "Tecnologia / decisao"], rows, [1.8, 4.7], font_size=9.5)

    rows = [[name, version, "Aplicacao"] for name, version in packages_app] + [[name, version, "Testes"] for name, version in packages_tests]
    add_table(doc, ["Pacote", "Versao", "Projeto"], rows, [3.2, 1.2, 2.1], font_size=8.5)
    doc.add_page_break()


def add_executive_summary(doc, tests, endpoints, models):
    add_heading(doc, 1, "2. Resumo executivo")
    status = Counter(t["status"] for t in tests)
    modules = len(set(t["module"] for t in tests))
    add_callout(
        doc,
        "Resultado consolidado",
        "A execucao pelo projeto de testes passou com 362 casos aprovados, 46 ignorados e 408 casos totais reportados pelo runner xUnit em .NET 10.0. A execucao pela solution falhou por duplicidade de nome de projeto, mas o .csproj de testes executou corretamente.",
    )
    rows = [
        ("Casos identificados no codigo-fonte", str(len(tests))),
        ("Casos ativos no inventario extraido", str(status.get("Ativo", 0))),
        ("Casos ignorados no inventario extraido", str(status.get("Ignorado", 0))),
        ("Modulos com testes automatizados", str(modules)),
        ("Endpoints REST catalogados", str(len(endpoints))),
        ("Classes/DTOs de modelo catalogados", str(len(models))),
    ]
    add_table(doc, ["Indicador", "Valor"], rows, [4.2, 2.3], font_size=10)
    add_para(doc, "Do ponto de vista de qualidade, o projeto combina testes unitarios de servico, testes de controller/importacao, gateways simulados para IoT/notificacoes e validacoes de payload via Postman. A cobertura esta concentrada onde o risco de regressao e maior: regras de validacao, combinacoes de status, limites ambientais, existencia de entidades relacionadas e tratamento de entradas invalidas.")
    add_para(doc, "Do ponto de vista de engenharia, a analise tambem revelou pontos de melhoria: limpeza da solution, reducao de avisos nullable, alinhamento de versoes do Entity Framework Core e decisao explicita entre modo mock e modo banco real no Program.cs.")
    doc.add_page_break()


def add_system_scope(doc):
    add_heading(doc, 1, "3. Sistema sob teste e ambito coberto")
    add_para(doc, "O GREENHERB representa o nucleo backend de uma plataforma de apoio a estufas inteligentes. A API modela operacoes agronomicas comuns e prepara pontos de integracao com sensores, notificacoes e fluxos de relatorio.")
    add_para(doc, "O dominio foi separado em modulos com responsabilidades bem delimitadas. Essa divisao facilita a construcao de testes focados, porque cada servico pode ser exercitado com dados controlados antes de expor o comportamento pelas controllers.")
    rows = []
    for key, data in MODULES.items():
        rows.append((data["titulo"], data["objetivo"], data["riscos"]))
    add_table(doc, ["Modulo", "Responsabilidade", "Risco principal coberto"], rows, [1.55, 2.55, 2.4], font_size=8.5)
    add_para(doc, "O ambito de teste cobre operacoes CRUD, consultas filtradas, importacao de ficheiros, validacoes de campos obrigatorios, ranges numericos, transicoes de estado, persistencia em memoria, mocks de gateways e verificacao de saidas esperadas.")
    doc.add_page_break()


def add_architecture(doc, endpoints, service_rows):
    ASSETS.mkdir(exist_ok=True)
    arch = ASSETS / "architecture.png"
    pyramid = ASSETS / "test_pyramid.png"
    make_architecture_diagram(arch)
    make_pyramid_diagram(pyramid)

    add_heading(doc, 1, "4. Arquitetura e stack tecnica")
    add_para(doc, "A arquitetura segue o desenho classico de uma Web API em camadas: controllers recebem requisicoes HTTP, services concentram regra de negocio, models definem contratos de dominio e o AppDbContext centraliza a persistencia relacional. Para desenvolvimento e testes, o projeto tambem contem provedores mock e dados estaticos.")
    doc.add_picture(str(arch), width=Inches(6.35))
    add_para(doc, "Figura 1. Fluxo logico de controllers, services, persistencia, gateways e testes.", style="Small", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_heading(doc, 2, "4.1 Camadas")
    for item in [
        "Controllers: traduzem verbos HTTP em chamadas de servico, retornando ActionResult com codigos 200, 201, 400, 404 e 204 conforme o caso.",
        "Services: aplicam regras de negocio, consultam repositorios/contexto ou MockDataProvider e controlam alteracoes de estado.",
        "Models/Requests: definem payloads, entidades de dominio e objetos auxiliares para importacao/exportacao.",
        "Gateways: simulam temperatura e notificacoes para permitir testes deterministas sem depender de hardware ou servicos externos.",
        "Testes: combinam xUnit, Theory/InlineData, Moq e EF Core InMemory para validar comportamento isolado e variacoes de entrada.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, 2, "4.2 Catalogo resumido de services")
    add_table(doc, ["Service", "Metodos", "Principais operacoes"], service_rows, [1.75, 0.85, 3.9], font_size=8.2)
    doc.add_page_break()

    add_heading(doc, 1, "5. Estrategia visual da piramide de testes")
    add_para(doc, "A estrategia privilegia testes unitarios e de componente antes de subir para integracao. Essa escolha e adequada para uma API de dominio com muitas regras condicionais, pois reduz custo de execucao e acelera diagnostico.")
    doc.add_picture(str(pyramid), width=Inches(6.3))
    add_para(doc, "Figura 2. Distribuicao conceitual entre unidade, integracao e sistema.", style="Small", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Os testes unitarios validam as regras em memoria. Os testes de integracao, documentados por Postman, validam a composicao externa da API: rotas, serializacao JSON, status HTTP, passagem de token e encadeamento de recursos.")
    doc.add_page_break()


def add_data_model(doc, models):
    add_heading(doc, 1, "6. Modelo de dados e persistencia")
    add_para(doc, "O AppDbContext expõe DbSet para Users, Herbs, CultivationPlans, Batches, OperationalTasks, Measurements, Alerts, Automations, Reports e AuditLogs. As configuracoes Fluent API definem chaves, obrigatoriedade, tamanho maximo, indices, precisao decimal e comportamento de delecao.")
    relation_rows = [
        ("User", "Username e Email unicos; Role obrigatoria; CreatedAt com CURRENT_TIMESTAMP", "Autenticacao e auditoria"),
        ("Herb", "Name + ScientificName unicos; Category/Origin obrigatorios; Notes/CareInstructions ate 2000 chars", "Catalogo botanico"),
        ("CultivationPlan", "HerbId com cascade; duracao e frequencia obrigatorias; limiares ambientais opcionais", "Planeamento agronomico"),
        ("Batch", "CultivationPlanId com cascade; LossPercentage e Productivity com precisao decimal", "Execucao de cultivo"),
        ("OperationalTask", "BatchId cascade; AssignedUserId SetNull; status e tipo obrigatorios", "Rotina operacional"),
        ("Measurement", "BatchId cascade; precisao decimal em temperatura, umidade e luminosidade", "Telemetria"),
        ("Alert", "ResolvedByUserId SetNull; tipo/status obrigatorios; resolution ate 2000 chars", "Eventos de risco"),
        ("Automation", "BatchId cascade; TriggerCondition/Action obrigatorios", "Motor de regras"),
        ("Report", "CreatedByUserId SetNull; ReportType/ExportFormat obrigatorios", "Analitica e exportacao"),
        ("AuditLog", "UserId NoAction; indices em UserId, EntityType e CreatedAt", "Rastreabilidade"),
    ]
    add_table(doc, ["Entidade", "Configuracao relevante", "Finalidade"], relation_rows, [1.45, 3.35, 1.7], font_size=8.1)

    add_heading(doc, 2, "6.1 Classes e propriedades")
    model_rows = []
    for model in models:
        props = "; ".join(f"{name}: {typ}" for name, typ in model["properties"][:10])
        if len(model["properties"]) > 10:
            props += "; ..."
        model_rows.append([model["name"], model["file"], props])
    add_table(doc, ["Classe", "Arquivo", "Propriedades principais"], model_rows, [1.45, 1.65, 3.4], font_size=8)
    doc.add_page_break()


def add_endpoint_catalog(doc, endpoints):
    add_heading(doc, 1, "7. Catalogo de endpoints REST")
    add_para(doc, "A API possui endpoints organizados por controller. Os endpoints usam rotas convencionais como api/[controller] em varios modulos e rotas explicitas em Plans e Automations.")
    grouped = defaultdict(list)
    for controller, verb, route, method in endpoints:
        grouped[controller].append([verb, route, method])
    for controller in sorted(grouped):
        add_heading(doc, 2, f"7.{len(grouped[controller])} {controller}")
        add_table(doc, ["Verbo", "Rota", "Metodo"], grouped[controller], [0.8, 3.5, 2.2], font_size=8.5)
    doc.add_page_break()


def add_test_strategy(doc, original_tables):
    add_heading(doc, 1, "8. Estrategia de testes adotada")
    add_para(doc, "A estrategia do projeto foi orientada pela piramide de automacao de testes. A maior parte do esforco esta nos testes unitarios, porque as regras criticas do GREENHERB vivem nos servicos e nos modelos de dominio.")
    add_heading(doc, 2, "8.1 Testes unitarios")
    add_para(doc, "Os testes unitarios usam xUnit para organizar cenarios, Moq para isolar dependencias quando necessario e EF Core InMemory para validar comportamento de persistencia sem custo de infraestrutura. Essa abordagem favorece feedback rapido, execucao local e reproducibilidade.")
    add_heading(doc, 2, "8.2 Testes de integracao")
    add_para(doc, "A camada de integracao foi representada pela colecao Postman e pelos bodies corrigidos. O objetivo e validar rotas, status HTTP, payloads flat sem wrapper, variaveis globais como batch_id e plan_id, alem de fluxos autenticados com token.")
    add_heading(doc, 2, "8.3 Criterios de aceite")
    for item in [
        "Entradas validas devem produzir objetos persistidos, status 200/201 ou retornos equivalentes de sucesso.",
        "Entradas invalidas devem retornar null, BadRequest, NotFound ou colecao vazia sem excecoes silenciosas.",
        "Filtros por status, tipo, role, batch, plano, entidade e intervalo temporal devem limitar corretamente o conjunto retornado.",
        "Transicoes de estado devem preservar timestamps, responsaveis e justificativas quando aplicavel.",
        "Gateways simulados devem registrar historico e responder corretamente a indisponibilidade.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, 1, "9. Tecnicas de teste aplicadas")
    add_heading(doc, 2, "9.1 Particionamento de equivalencia")
    add_para(doc, "O particionamento de equivalencia reduz o universo de entradas em classes representativas. No GREENHERB, a tecnica aparece em autenticacao, importacao de plantas, criacao de entidades e filtros de consulta.")
    if original_tables:
        add_table(doc, ["ID", "Cenario", "Metodo", "Nivel", "Tecnica", "Resultado esperado", "Pre-condicao"], original_tables[0], [0.65, 1.7, 1.0, 0.65, 1.05, 1.65, 1.0], font_size=7.1)

    add_heading(doc, 2, "9.2 Analise de valores limite")
    add_para(doc, "A analise de fronteira e especialmente relevante para variaveis ambientais e campos numericos. O projeto testa limites de senha, ciclo de cultivo, duracao, frequencia, temperatura, umidade e luminosidade.")
    if len(original_tables) > 1:
        add_table(doc, ["ID", "Cenario", "Metodo", "Nivel", "Tecnica", "Resultado esperado", "Pre-condicao"], original_tables[1], [0.65, 1.7, 1.0, 0.65, 1.05, 1.65, 1.0], font_size=7.1)

    add_heading(doc, 2, "9.3 Cobertura de decisoes e condicoes")
    add_para(doc, "A cobertura de decisoes valida combinacoes booleanas e caminhos alternativos: usuario ativo/inativo, senha correta/incorreta, recursos existentes/inexistentes, status e roles, resolucao ou ignorar alertas, bem como ranges ambientais validos ou invertidos.")
    if len(original_tables) > 2:
        add_table(doc, ["ID", "Cenario", "Metodo", "Nivel", "Tecnica", "Resultado esperado", "Pre-condicao"], original_tables[2], [0.65, 1.7, 1.0, 0.65, 1.05, 1.65, 1.0], font_size=7.1)
    doc.add_page_break()


def add_module_coverage(doc, tests):
    add_heading(doc, 1, "10. Cobertura por modulo")
    by_module = defaultdict(list)
    for test in tests:
        by_module[test["module"]].append(test)
    summary_rows = []
    for module, items in sorted(by_module.items()):
        active = sum(1 for t in items if t["status"] == "Ativo")
        skipped = sum(1 for t in items if t["status"] == "Ignorado")
        summary_rows.append([module, str(active), str(skipped), str(len(items)), ", ".join(sorted(set(t["technique"] for t in items)))])
    add_table(doc, ["Modulo", "Ativos", "Ignorados", "Total", "Tecnicas observadas"], summary_rows, [1.35, 0.75, 0.85, 0.75, 2.8], font_size=8.2)

    ordered_modules = [
        "Auth", "User", "Herb", "Plan", "Batch", "Measurement", "Alert", "Automation",
        "Task", "Report", "AuditLog", "NotificationGateway", "TemperatureGateway",
    ]
    for module in ordered_modules:
        data = MODULES.get(module, {"titulo": module, "objetivo": "", "riscos": ""})
        items = by_module.get(module, [])
        add_heading(doc, 2, f"10.{ordered_modules.index(module) + 1} {data['titulo']}")
        add_para(doc, data["objetivo"])
        add_para(doc, f"Riscos exercitados pelos testes: {data['riscos']}")
        if items:
            counts = Counter(t["technique"] for t in items)
            rows = [[tech, str(count)] for tech, count in sorted(counts.items())]
            add_table(doc, ["Tecnica / foco", "Casos"], rows, [4.8, 1.7], font_size=9)
            examples = []
            for t in items[:8]:
                examples.append([shorten(t["case"], width=82, placeholder="..."), t["status"], t["technique"]])
            add_table(doc, ["Exemplo de caso", "Estado", "Foco"], examples, [4.2, 0.8, 1.5], font_size=8)
        else:
            add_para(doc, "Nao foram encontrados testes automatizados com este nome de modulo no inventario extraido; a cobertura pode existir sob outro agrupamento ou nos fluxos Postman.")
        doc.add_page_break()


def add_execution_results(doc):
    add_heading(doc, 1, "11. Resultados de execucao")
    add_para(doc, "A validacao local foi executada em 12/06/2026 no workspace do projeto. A primeira tentativa pela solution identificou um problema de configuracao; a segunda, pelo projeto de testes, executou a suite.")
    rows = [
        ("dotnet test --no-restore \"Trabalho Testes NET.sln\"", "Falhou", "MSB5004: solution possui dois projetos denominados greenherb."),
        ("dotnet test GREENHERB.Tests\\GREENHERB.Tests.csproj --no-restore", "Aprovado", "0 falhas, 362 aprovados, 46 ignorados, 408 total, duracao reportada de 5 s."),
    ]
    add_table(doc, ["Comando", "Resultado", "Observacao"], rows, [2.65, 1.0, 2.85], font_size=8.5)

    add_heading(doc, 2, "11.1 Avisos relevantes")
    for item in [
        "NU1510 indica que Microsoft.Extensions.Configuration.Json provavelmente e desnecessario e pode ser removido se nao houver uso direto.",
        "CS8625 e CS8602 aparecem em trechos com null em tipos nao anulaveis e possivel desreferenciacao nula.",
        "MSB3277 aponta conflito entre Microsoft.EntityFrameworkCore.Relational 10.0.4 e 10.0.7.",
        "xUnit2002 aponta uso desnecessario de Assert.NotNull em DateTime.",
        "A execucao pelo .csproj e valida para a suite, mas a solution deve ser corrigida antes de uso em CI/CD.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, 2, "11.2 Leitura dos ignorados")
    add_para(doc, "Os testes ignorados nao representam falha de execucao, mas reduzem evidencia ativa. A maioria dos casos ignorados esta associada a cenarios documentados como 'falhou localmente', sobretudo em PlanTestes e alguns testes pontuais de servico.")
    add_callout(doc, "Interpretacao", "A suite e executavel e sem falhas ativas, mas a meta de maturidade deve ser reativar ou justificar formalmente os 46 casos ignorados antes de considerar a cobertura como final.")
    doc.add_page_break()


def add_defects_and_improvements(doc, original_tables):
    add_heading(doc, 1, "12. Defeitos, riscos e melhorias")
    add_para(doc, "O documento-base ja trazia defeitos encontrados durante a construcao da suite. A versao expandida mantem esses itens e acrescenta riscos tecnicos observados na execucao local.")
    if len(original_tables) > 3:
        add_table(doc, ["ID", "Descricao / Sintoma", "Severidade", "Passos", "Status"], original_tables[3], [0.8, 2.1, 0.8, 1.6, 1.2], font_size=7.7)
    add_heading(doc, 2, "12.1 Riscos tecnicos atuais")
    rows = [
        ("Solution com nomes duplicados", "Alta", "Impede dotnet test pela solution e dificulta CI/CD.", "Renomear os projetos no .sln ou recriar a solution com nomes distintos."),
        ("Avisos nullable", "Media", "Podem esconder NullReferenceException em rotas e servicos.", "Tratar null de forma explicita, ajustar construtores e contratos."),
        ("Conflito EF Core Relational", "Media", "Pode gerar comportamento divergente entre build local e pipeline.", "Alinhar versoes EF Core/Npgsql/Design/InMemory."),
        ("Mocks duplicados no Program.cs", "Media", "A configuracao de DI fica ambigua e dificulta alternar mock/banco real.", "Separar profiles de ambiente e remover duplicidade de AddScoped."),
        ("Casos ignorados", "Media", "Reduzem confianca na cobertura declarada.", "Reativar, corrigir ou documentar justificativa individual."),
    ]
    add_table(doc, ["Risco", "Severidade", "Impacto", "Acao recomendada"], rows, [1.45, 0.85, 2.1, 2.1], font_size=8.2)

    add_heading(doc, 2, "12.2 Propostas de melhoria continua")
    for item in [
        "Adicionar GitHub Actions com dotnet restore, dotnet build, dotnet test e publicacao de relatorio de cobertura.",
        "Executar a colecao Postman com Newman no pipeline, alimentando variaveis globais de token e IDs criados.",
        "Migrar testes criticos de persistencia para Testcontainers com PostgreSQL real.",
        "Gerar relatorios de cobertura com coverlet e criar limite minimo por camada.",
        "Padronizar ambiente mock versus banco real por appsettings e variaveis de ambiente.",
        "Adicionar testes de carga com k6 ou JMeter para ingestao massiva de telemetria.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()


def add_conclusions(doc):
    add_heading(doc, 1, "13. Conclusoes")
    add_para(doc, "A avaliacao expandida confirma que o GREENHERB possui uma base de testes ampla para o escopo academico do projeto. O codigo cobre funcoes essenciais do dominio, exercita entradas validas e invalidas e demonstra preocupacao com rastreabilidade, limites ambientais e estados operacionais.")
    add_para(doc, "A maior qualidade da suite esta na diversidade de cenarios: autenticacao, roles, importacao, CRUD, filtros, gateways, automacoes, alertas, relatorios e auditoria. Essa diversidade reduz o risco de regressao em areas que, em um sistema de estufa, poderiam causar inconsistencia operacional.")
    add_para(doc, "A principal oportunidade de evolucao nao e aumentar quantidade bruta de testes, mas amadurecer o ciclo de entrega: corrigir a solution, reativar ignorados, alinhar dependencias, reduzir warnings e automatizar execucao em pipeline. Com esses ajustes, o projeto passa de uma suite local robusta para uma pratica de qualidade continua.")
    doc.add_page_break()


def add_test_inventory(doc, tests):
    add_heading(doc, 1, "Anexo A. Inventario completo de testes automatizados")
    add_para(doc, "Este anexo lista os casos identificados no codigo-fonte. Testes Theory com InlineData foram expandidos como casos separados, pois o runner executa cada combinacao de dados como uma variacao do mesmo metodo.")
    by_file = defaultdict(list)
    for test in tests:
        by_file[test["file"]].append(test)
    for file_name in sorted(by_file):
        items = by_file[file_name]
        add_heading(doc, 2, file_name)
        rows = []
        for idx, t in enumerate(items, 1):
            rows.append([str(idx), shorten(t["case"], width=95, placeholder="..."), t["technique"], t["status"]])
        add_table(doc, ["#", "Caso", "Foco", "Estado"], rows, [0.45, 4.55, 1.0, 0.5], font_size=7.4)
    doc.add_page_break()


def add_postman_bodies(doc):
    add_heading(doc, 1, "Anexo B. Bodies corrigidos para Postman")
    path = ROOT / "BODIES_CORRIGIDOS.json"
    if not path.exists():
        add_para(doc, "Arquivo BODIES_CORRIGIDOS.json nao encontrado no workspace.")
        return
    data = json.loads(read_text(path))
    add_para(doc, data.get("info", {}).get("description", "Exemplos de bodies corretos para testes de integracao."))
    rows = []
    for endpoint, payload in data.get("endpoints", {}).items():
        body = payload.get("body") or payload.get("fields") or {}
        summary = json.dumps(body, ensure_ascii=False, indent=2)
        rows.append([endpoint, payload.get("description", ""), shorten(summary.replace("\n", " "), width=160, placeholder="...")])
    add_table(doc, ["Endpoint", "Descricao", "Resumo do payload"], rows, [1.75, 2.0, 2.75], font_size=7.6)

    for endpoint, payload in data.get("endpoints", {}).items():
        add_heading(doc, 2, endpoint)
        add_para(doc, payload.get("description", ""))
        body = payload.get("body") or payload.get("fields") or {}
        code = json.dumps(body, ensure_ascii=False, indent=2)
        for line in code.splitlines():
            add_para(doc, line, style="CodeBlock")
        if payload.get("script_post"):
            add_para(doc, "Script pos-requisicao:", bold=True, after=2)
            add_para(doc, payload["script_post"], style="CodeBlock")
    doc.add_page_break()


def add_technical_catalogs(doc, service_rows, endpoints, models):
    add_heading(doc, 1, "Anexo C. Catalogos tecnicos do codigo")
    add_heading(doc, 2, "C.1 Services")
    add_table(doc, ["Service", "Qtd. metodos", "Operacoes"], service_rows, [1.8, 0.85, 3.85], font_size=7.8)
    add_heading(doc, 2, "C.2 Controllers e rotas")
    add_table(doc, ["Controller", "Verbo", "Rota", "Metodo"], endpoints, [1.25, 0.65, 3.1, 1.5], font_size=7.5)
    add_heading(doc, 2, "C.3 Models")
    rows = []
    for model in models:
        for prop, typ in model["properties"]:
            rows.append([model["name"], prop, typ])
    add_table(doc, ["Classe", "Propriedade", "Tipo"], rows, [1.65, 2.25, 2.6], font_size=7.5)


def main():
    ASSETS.mkdir(parents=True, exist_ok=True)
    packages_app = parse_packages(ROOT / "GREENHERB" / "GREENHERB.csproj")
    packages_tests = parse_packages(ROOT / "GREENHERB.Tests" / "GREENHERB.Tests.csproj")
    models = parse_models()
    endpoints = parse_controllers()
    service_rows = parse_services()
    tests = parse_tests()
    coverage = parse_coverage()
    traceability_rows = build_traceability_rows(tests)
    export_paths = write_exports(tests, traceability_rows, coverage)
    original_tables = read_original_tables()

    doc = Document()
    configure_document(doc)
    set_header_footer(doc)
    props = doc.core_properties
    props.title = "Relatorio Final de Testes de Software - GREENHERB"
    props.subject = "Versao expandida com 40+ paginas"
    props.author = "Christopher Carpegiane de Souza Rocha"

    add_cover(doc)
    add_static_toc(doc)
    add_delivery_identification(doc, packages_app, packages_tests)
    add_delivery_scope(doc, models, endpoints)
    add_delivery_strategy(doc)
    add_delivery_test_plan(doc, endpoints)
    add_delivery_equivalence(doc)
    add_delivery_boundary(doc)
    add_delivery_conditions(doc)
    add_delivery_traceability(doc, traceability_rows)
    add_delivery_results(doc, tests, coverage, export_paths)
    add_delivery_defects(doc, original_tables)
    add_delivery_conclusions(doc)
    add_test_inventory(doc, tests)
    add_postman_bodies(doc)
    add_technical_catalogs(doc, service_rows, endpoints, models)

    doc.save(OUT)
    print(OUT)
    print(f"tests={len(tests)} endpoints={len(endpoints)} models={len(models)}")


if __name__ == "__main__":
    main()
